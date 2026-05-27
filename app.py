import asyncio
import csv
import json
import os
from io import BytesIO
from io import StringIO

import pandas as pd
import streamlit as st

from agent import ALLOWED_INTENTS, MODEL_PROVIDER, orchestrate_request
from evaluator import run_benchmark

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:
    A4 = None
    canvas = None

try:
    from docx import Document
except Exception:
    Document = None


EXPECTED_COLUMNS = ["input_text", "actual_label"]


def normalize_label(label: str) -> str:
    normalized = (label or "").strip().lower()
    _alias_map = {
        "greeting":        "greetings",
        "greetings":       "greetings",
        "device_control":  "device_control",
        "service_request": "service_request",
        "service":         "service_request",
        "shopping":        "service_request",
        "queries":         "service_request",
        "automations":     "automations",
        "automation":      "automations",
        "out_of_scope":    "out_of_scope",
        "guardrail":       "out_of_scope",
        "unsafe":          "out_of_scope",
    }
    if normalized in _alias_map:
        return _alias_map[normalized]
    for key, canonical in _alias_map.items():
        if normalized.startswith(key):
            return canonical
    return "out_of_scope"


def run_orchestrator(text: str):
    return asyncio.run(orchestrate_request(text))


def dataframe_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def _build_summary_lines(report: dict) -> list[str]:
    lines = []
    lines.append("Orchestrator Benchmark Summary")
    lines.append("")
    lines.append(f"Model Provider: {report['model_config']['model_provider']}")
    lines.append(f"Model ID: {report['model_config']['model_id']}")
    lines.append(f"Dataset: {report['model_config']['dataset']}")
    lines.append(f"Total Requests: {report['model_config']['total_requests']}")
    lines.append("")
    lines.append("Latency (seconds)")
    lines.append(f"P50: {report['latency_seconds']['p50']}")
    lines.append(f"P90: {report['latency_seconds']['p90']}")
    lines.append(f"P99: {report['latency_seconds']['p99']}")
    lines.append("")
    lines.append("Classification")
    lines.append(f"Overall Accuracy: {report['classification']['overall_accuracy']}")
    lines.append(f"Guardrail False Positive Rate: {report['guardrail']['false_positive_rate']}")
    lines.append("")
    lines.append("Token Usage")
    lines.append(f"Avg Input Tokens: {report['token_usage']['avg_input_tokens']}")
    lines.append(f"Avg Output Tokens: {report['token_usage']['avg_output_tokens']}")
    lines.append("")
    lines.append("Cost")
    lines.append(f"Total Cost: {report['cost']['total_cost']}")
    lines.append(f"Avg Cost/Request: {report['cost']['avg_cost_per_request']}")
    lines.append(f"Cost per 1K Queries: {report['cost']['cost_per_1k_queries']}")
    lines.append("")
    lines.append("Quality")
    lines.append(f"JSON Validity Rate: {report['quality']['json_validity_rate']}")
    lines.append(f"Schema Compliance Rate: {report['quality']['schema_compliance_rate']}")
    lines.append(f"Failure Count: {report['quality']['failure_count']}")
    lines.append(f"Retry Count: {report['quality']['retry_count']}")
    lines.append("")
    lines.append("Per-Intent Metrics")
    for intent, vals in report["classification"]["per_intent"].items():
        lines.append(
            f"{intent}: precision={vals.get('precision', 0)}, recall={vals.get('recall', 0)}, f1={vals.get('f1', 0)}"
        )

    lines.append("")
    lines.append("Multilingual Accuracy")
    for lang, vals in report["multilingual"].items():
        lines.append(f"{lang}: accuracy={vals.get('accuracy', 0)}, correct={vals.get('correct', 0)}, total={vals.get('total', 0)}")

    lines.append("")
    lines.append("Misclassification Matrix")
    matrix = report.get("misclassification_matrix", {})
    labels = list(matrix.keys())
    if labels:
        lines.append("actual/predicted," + ",".join(labels))
        for actual in labels:
            row = [str(matrix.get(actual, {}).get(pred, 0)) for pred in labels]
            lines.append(f"{actual}," + ",".join(row))

    return lines


def build_pdf_report(report: dict) -> bytes | None:
    if canvas is None or A4 is None:
        return None

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 40
    lines = _build_summary_lines(report)

    for line in lines:
        if y < 40:
            pdf.showPage()
            y = height - 40
        pdf.drawString(40, y, str(line)[:150])
        y -= 14

    pdf.save()
    buffer.seek(0)
    return buffer.read()


def build_detailed_pdf_report(report: dict, detailed_df: pd.DataFrame, max_rows: int = 200) -> bytes | None:
    if canvas is None or A4 is None:
        return None

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4
    y = height - 40

    for line in _build_summary_lines(report):
        if y < 40:
            pdf.showPage()
            y = height - 40
        pdf.drawString(40, y, str(line)[:150])
        y -= 14

    if y < 70:
        pdf.showPage()
        y = height - 40
    pdf.drawString(40, y, "Detailed Results Preview")
    y -= 18

    preview = detailed_df.head(max_rows)
    headers = [str(col) for col in preview.columns]
    header_line = " | ".join(headers)
    if y < 40:
        pdf.showPage()
        y = height - 40
    pdf.drawString(40, y, header_line[:150])
    y -= 14

    for _, row in preview.iterrows():
        if y < 40:
            pdf.showPage()
            y = height - 40
        row_line = " | ".join(str(row[col]) for col in preview.columns)
        pdf.drawString(40, y, row_line[:150])
        y -= 12

    pdf.save()
    buffer.seek(0)
    return buffer.read()


def build_docx_report(report: dict, detailed_df: pd.DataFrame) -> bytes | None:
    if Document is None:
        return None

    doc = Document()
    doc.add_heading("Orchestrator Benchmark Report", level=1)

    for line in _build_summary_lines(report):
        doc.add_paragraph(line)

    doc.add_heading("Detailed Results (Top 200 rows)", level=2)
    preview = detailed_df.head(200)
    table = doc.add_table(rows=1, cols=len(preview.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(preview.columns):
        hdr_cells[i].text = str(col)

    for _, row in preview.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(preview.columns):
            cells[i].text = str(row[col])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def build_summary_docx_report(report: dict) -> bytes | None:
    if Document is None:
        return None

    doc = Document()
    doc.add_heading("Orchestrator Benchmark Summary", level=1)

    for line in _build_summary_lines(report):
        doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


st.set_page_config(page_title="Orchestrator Agent UI", page_icon="AI", layout="wide")

st.title("Orchestrator Agent - Web UI")
st.caption(f"Model provider: {MODEL_PROVIDER} | Intents: {', '.join(ALLOWED_INTENTS)}")

tab_live, tab_eval, tab_schema = st.tabs(["Live Query", "Dataset Evaluation", "Dataset Format"])


with tab_live:
    st.subheader("Test Single User Query")
    prompt = st.text_area("User input", placeholder="Example: Turn on AC in bedroom", height=120)
    if st.button("Run Orchestrator", type="primary"):
        if not prompt.strip():
            st.warning("Please enter a query.")
        else:
            with st.spinner("Running orchestrator..."):
                try:
                    import asyncio as _asyncio
                    from agent import orchestrate_request_with_meta as _orch_meta
                    full_result = _asyncio.run(_orch_meta(prompt))

                    mixed_out = full_result.get("output", {})
                    multi_out = full_result.get("multi_intent_output", {})

                    st.success("Done")

                    col_mixed, col_multi = st.columns(2)

                    with col_mixed:
                        st.markdown(
                            "**🔵 Mixed-Intent Result**  "
                            "*(single best label — production default)*"
                        )
                        st.caption(
                            "Trained on: final\_cleaned\_dataset + new\_dataset + "
                            "havells\_multi\_intent\_dataset  "
                            "| Each query → **one** canonical intent label."
                        )
                        st.json(mixed_out)

                    with col_multi:
                        st.markdown(
                            "**🟠 Multi-Intent Result**  "
                            "*(two labels + priority rules applied)*"
                        )
                        st.caption(
                            "Trained on: havells\_multi\_intent\_dataset only  "
                            "| Each query → **one or two** intent labels with business rules."
                        )
                        if multi_out:
                            st.json(multi_out)
                        else:
                            st.info(
                                "Multi-intent model not loaded. "
                                "Run `python train_multi_intent_classifier.py` first."
                            )

                except Exception as exc:
                    st.error(f"Failed: {exc}")


with tab_eval:
    st.subheader("Upload Dataset and Run Full Benchmark")
    st.write("Accepted formats:  **CSV** (columns: `input_text`, `actual_label`)  |  **JSON** (fields: `query`, `expected_response_type`)")

    # ── Mode selector ──────────────────────────────────────────────────
    eval_mode = st.radio(
        "Evaluation Mode",
        options=["Mixed-Intent", "Multi-Intent"],
        horizontal=True,
        help=(
            "**Mixed-Intent**: evaluates using the single-label classifier "
            "(tier2_classifier.pkl). Best for: final_cleaned_dataset.csv, new_dataset.csv.\n\n"
            "**Multi-Intent**: evaluates using the multi-label classifier "
            "(tier2_multi_intent_classifier.pkl) with business priority rules. "
            "Best for: havells_multi_intent_dataset.csv."
        ),
    )
    mode_key = "mixed" if eval_mode == "Mixed-Intent" else "multi"

    if mode_key == "multi":
        st.info(
            "🟠 **Multi-Intent mode selected.** "
            "The benchmark will use the dedicated multi-intent classifier output. "
            "Accuracy is computed as exact set-match (both labels must match). "
            "Make sure you have run `python train_multi_intent_classifier.py` first."
        )
    else:
        st.info(
            "🔵 **Mixed-Intent mode selected.** "
            "The benchmark will use the standard single-label classifier output."
        )

    cost_col1, cost_col2 = st.columns(2)
    with cost_col1:
        input_cost = st.number_input("Input cost per 1K tokens", min_value=0.0, value=0.0, step=0.0001, format="%.6f")
    with cost_col2:
        output_cost = st.number_input("Output cost per 1K tokens", min_value=0.0, value=0.0, step=0.0001, format="%.6f")

    uploaded = st.file_uploader(
        "Upload dataset (CSV or JSON)",
        type=["csv", "json"],
        help=(
            "**CSV**: must have columns `input_text` and `actual_label`.\n\n"
            "**JSON**: array of objects with fields `query` and `expected_response_type`. "
            "Optionally include `language` for multilingual breakdown."
        ),
    )

    if uploaded is not None:
        file_ext = uploaded.name.rsplit(".", 1)[-1].lower()
        rows = []
        load_error = None

        # ── JSON branch ──────────────────────────────────────────────────────
        if file_ext == "json":
            try:
                raw_json = json.loads(uploaded.read().decode("utf-8", errors="replace"))
                if not isinstance(raw_json, list):
                    load_error = "JSON file must contain a top-level array of objects."
                else:
                    for obj in raw_json:
                        query = str(obj.get("query", "")).strip()
                        label = str(obj.get("expected_response_type", "")).strip()
                        if not query or not label:
                            continue
                        row = {"input_text": query, "actual_label": label}
                        # Carry language through for multilingual breakdown
                        if "language" in obj:
                            row["language"] = obj["language"]
                        rows.append(row)
                    if not rows:
                        load_error = "No valid records found. Each object needs `query` and `expected_response_type`."
                    else:
                        st.success(
                            f"JSON loaded: **{len(rows):,} records** — "
                            f"mapped `query` → `input_text`, `expected_response_type` → `actual_label`."
                        )
            except Exception as exc:
                load_error = f"Failed to parse JSON: {exc}"

        # ── CSV branch ───────────────────────────────────────────────────────
        else:
            content = uploaded.read().decode("utf-8", errors="replace")
            reader = csv.DictReader(StringIO(content))
            headers = reader.fieldnames or []
            missing = [col for col in EXPECTED_COLUMNS if col not in headers]
            if missing:
                load_error = f"Missing required columns: {', '.join(missing)}"
            else:
                rows = list(reader)

        if load_error:
            st.error(load_error)
        elif rows:
            st.info(f"Loaded **{len(rows):,} rows** — will evaluate in **{eval_mode}** mode.")

            if st.button("Run Full Evaluation", type="primary"):
                with st.spinner(f"Running {eval_mode} benchmark pipeline..."):
                    output_dir = "benchmark_outputs"
                    os.makedirs(output_dir, exist_ok=True)

                    dataset_path = os.path.join(output_dir, "uploaded_dataset.csv")
                    csv_headers = list(rows[0].keys())
                    with open(dataset_path, "w", encoding="utf-8", newline="") as f:
                        writer = csv.DictWriter(f, fieldnames=csv_headers)
                        writer.writeheader()
                        writer.writerows(rows)

                    old_input_cost = os.environ.get("INPUT_COST_PER_1K_TOKENS")
                    old_output_cost = os.environ.get("OUTPUT_COST_PER_1K_TOKENS")
                    os.environ["INPUT_COST_PER_1K_TOKENS"] = str(input_cost)
                    os.environ["OUTPUT_COST_PER_1K_TOKENS"] = str(output_cost)

                    try:
                        result = asyncio.run(run_benchmark(dataset_path, output_dir, mode=mode_key))
                    finally:
                        if old_input_cost is None:
                            os.environ.pop("INPUT_COST_PER_1K_TOKENS", None)
                        else:
                            os.environ["INPUT_COST_PER_1K_TOKENS"] = old_input_cost

                        if old_output_cost is None:
                            os.environ.pop("OUTPUT_COST_PER_1K_TOKENS", None)
                        else:
                            os.environ["OUTPUT_COST_PER_1K_TOKENS"] = old_output_cost

                report = result["report"]
                paths = result["paths"]

                metrics_row_1 = st.columns(4)
                metrics_row_1[0].metric("Total Requests", report["model_config"]["total_requests"])
                metrics_row_1[1].metric("Accuracy", f"{report['classification']['overall_accuracy'] * 100:.2f}%")
                metrics_row_1[2].metric("Guardrail FPR", f"{report['guardrail']['false_positive_rate'] * 100:.2f}%")
                metrics_row_1[3].metric("JSON Validity", f"{report['quality']['json_validity_rate'] * 100:.2f}%")
                st.caption(f"Mode: **{report['model_config'].get('evaluation_mode', 'mixed').upper()}**")

                metrics_row_2 = st.columns(4)
                metrics_row_2[0].metric("Latency P50 (s)", report["latency_seconds"]["p50"])
                metrics_row_2[1].metric("Latency P90 (s)", report["latency_seconds"]["p90"])
                metrics_row_2[2].metric("Latency P99 (s)", report["latency_seconds"]["p99"])
                metrics_row_2[3].metric("Schema Compliance", f"{report['quality']['schema_compliance_rate'] * 100:.2f}%")

                st.subheader("Token and Cost")
                token_cost_df = pd.DataFrame(
                    [
                        {
                            "avg_input_tokens": report["token_usage"]["avg_input_tokens"],
                            "avg_output_tokens": report["token_usage"]["avg_output_tokens"],
                            "total_cost": report["cost"]["total_cost"],
                            "avg_cost_per_request": report["cost"]["avg_cost_per_request"],
                            "cost_per_1k_queries": report["cost"]["cost_per_1k_queries"],
                            "failure_count": report["quality"]["failure_count"],
                            "retry_count": report["quality"]["retry_count"],
                        }
                    ]
                )
                st.dataframe(token_cost_df, use_container_width=True)

                st.subheader("Per-Intent Precision / Recall / F1")
                per_intent = report["classification"]["per_intent"]
                per_intent_df = pd.DataFrame(per_intent).T.reset_index().rename(columns={"index": "intent"})
                st.dataframe(per_intent_df, use_container_width=True)

                st.subheader("Multilingual Accuracy")
                multilingual_df = pd.DataFrame(report["multilingual"]).T.reset_index().rename(columns={"index": "language"})
                st.dataframe(multilingual_df, use_container_width=True)

                st.subheader("Misclassification Matrix")
                matrix_df = pd.DataFrame(report["misclassification_matrix"]).T
                matrix_df.index.name = "actual \\ predicted"
                st.dataframe(matrix_df, use_container_width=True)

                st.subheader("Detailed Results")
                with open(paths["detailed_csv"], "r", encoding="utf-8") as f:
                    detailed_df = pd.read_csv(f)
                st.dataframe(detailed_df, use_container_width=True)

                st.subheader("Download Reports")
                with open(paths["json_report"], "r", encoding="utf-8") as f:
                    json_payload = f.read()
                with open(paths["summary_csv"], "r", encoding="utf-8") as f:
                    summary_payload = f.read()
                with open(paths["detailed_csv"], "r", encoding="utf-8") as f:
                    detailed_payload = f.read()
                with open(paths["misclassification_matrix_csv"], "r", encoding="utf-8") as f:
                    matrix_payload = f.read()

                st.download_button(
                    label="Download benchmark_report.json",
                    data=json_payload,
                    file_name="benchmark_report.json",
                    mime="application/json",
                )
                st.download_button(
                    label="Download benchmark_summary.csv",
                    data=summary_payload,
                    file_name="benchmark_summary.csv",
                    mime="text/csv",
                )
                st.download_button(
                    label="Download benchmark_detailed_results.csv",
                    data=detailed_payload,
                    file_name="benchmark_detailed_results.csv",
                    mime="text/csv",
                )
                st.download_button(
                    label="Download misclassification_matrix.csv",
                    data=matrix_payload,
                    file_name="misclassification_matrix.csv",
                    mime="text/csv",
                )

                pdf_bytes = build_pdf_report(report)
                if pdf_bytes is not None:
                    st.download_button(
                        label="Download summary PDF",
                        data=pdf_bytes,
                        file_name="benchmark_summary_report.pdf",
                        mime="application/pdf",
                    )
                else:
                    st.info("Install reportlab to enable PDF export: pip install reportlab")

                docx_bytes = build_docx_report(report, detailed_df)
                if docx_bytes is not None:
                    st.download_button(
                        label="Download detailed DOCX",
                        data=docx_bytes,
                        file_name="benchmark_detailed_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                else:
                    st.info("Install python-docx to enable DOCX export: pip install python-docx")

                summary_docx_bytes = build_summary_docx_report(report)
                if summary_docx_bytes is not None:
                    st.download_button(
                        label="Download summary DOCX",
                        data=summary_docx_bytes,
                        file_name="benchmark_summary_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                detailed_pdf_bytes = build_detailed_pdf_report(report, detailed_df)
                if detailed_pdf_bytes is not None:
                    st.download_button(
                        label="Download detailed PDF",
                        data=detailed_pdf_bytes,
                        file_name="benchmark_detailed_report.pdf",
                        mime="application/pdf",
                    )

                st.expander("Raw JSON report").json(json.loads(json_payload))


with tab_schema:
    st.subheader("Expected Dataset Schema")

    st.markdown("### Option 1 — JSON format (recommended for `final_dataset_improved_automation.json`)")
    st.markdown(
        """
        - File type: **JSON** (top-level array of objects)
        - Required fields per object:
          - **`query`** — user query text
          - **`expected_response_type`** — ground truth intent label
        - Optional fields (carried through automatically):
          - **`language`** — used for per-language accuracy breakdown
          - `id`, `is_boundary`, `sub_category` — ignored
        - Supported intent values: `automation` / `automations`, `device_control`, `greetings`, `out_of_scope`, `service_request`
        """
    )
    st.code(
        '[\n'
        '  {"query": "Turn on the fan", "expected_response_type": "device_control", "language": "english"},\n'
        '  {"query": "AC ko band karo", "expected_response_type": "device_control", "language": "hindi"},\n'
        '  {"query": "Hello good morning", "expected_response_type": "greetings",      "language": "english"}\n'
        ']',
        language="json",
    )

    st.markdown("### Option 2 — CSV format")
    st.markdown(
        """
        - File type: **CSV**
        - Required columns:
          - **`input_text`** — user query text
          - **`actual_label`** — ground truth intent label(s)
        - **Single-label** (one label per row): `greetings`, `device_control`, `service_request`, `automations`, `out_of_scope`
        - **Multi-label** (comma-separated): `device_control,service_request` or `automations,service_request`
        - Backward-compatible aliases:
          - `greeting` → `greetings`
          - `service` → `service_request`
          - `automation` → `automations`
          - `shopping` / `queries` → `service_request`
          - `guardrail` → `out_of_scope`
        """
    )

    st.markdown("#### Mixed-Intent Sample")
    sample_mixed_df = pd.DataFrame(
        [
            {"input_text": "Hi there", "actual_label": "greetings"},
            {"input_text": "Turn on bedroom fan", "actual_label": "device_control"},
            {"input_text": "Turn on AC at 9pm daily", "actual_label": "automations"},
            {"input_text": "Explore AMC plans for AC", "actual_label": "service_request"},
            {"input_text": "Hack my neighbor wifi", "actual_label": "out_of_scope"},
        ]
    )
    st.dataframe(sample_mixed_df, use_container_width=True)
    st.download_button(
        label="Download mixed-intent sample CSV",
        data=dataframe_to_csv_bytes(sample_mixed_df),
        file_name="sample_mixed_intent_dataset.csv",
        mime="text/csv",
    )

    st.markdown("#### Multi-Intent Sample")
    sample_multi_df = pd.DataFrame(
        [
            {"input_text": "Turn on the AC and raise a service request for it", "actual_label": "device_control,service_request"},
            {"input_text": "Switch on the fan and explore AMC plans for it", "actual_label": "device_control,service_request"},
            {"input_text": "Turn on AC at 9pm and explore AMC plans", "actual_label": "automations,service_request"},
            {"input_text": "Schedule geyser for 5am and register my geyser", "actual_label": "automations,service_request"},
            {"input_text": "Hello, raise a complaint about my AC", "actual_label": "service_request"},
        ]
    )
    st.dataframe(sample_multi_df, use_container_width=True)
    st.download_button(
        label="Download multi-intent sample CSV",
        data=dataframe_to_csv_bytes(sample_multi_df),
        file_name="sample_multi_intent_dataset.csv",
        mime="text/csv",
    )